import os
from pathlib import Path
from typing import List
import pandas as pd
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader


class DocumentIngestor:
    """Multi-format document loader supporting PDF, DOCX, TXT, and Excel/CSV files."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".xlsx", ".xls", ".csv"}

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load_file(self, file_path: str) -> List[Document]:
        """Load a single document based on its file extension."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format '{ext}'. Supported: {self.SUPPORTED_EXTENSIONS}")

        if ext == ".pdf":
            return self._load_pdf(path)
        elif ext in {".docx", ".doc"}:
            return self._load_docx(path)
        elif ext in {".txt", ".md"}:
            return self._load_txt(path)
        elif ext in {".xlsx", ".xls", ".csv"}:
            return self._load_excel_csv(path)
        else:
            raise ValueError(f"No handler configured for file format: {ext}")

    def _load_pdf(self, path: Path) -> List[Document]:
        """Load PDF document pages with dual parser fallback (PyPDFLoader + pdfplumber)."""
        docs = []
        try:
            loader = PyPDFLoader(str(path))
            docs = loader.load()
        except Exception as e:
            print(f"[Warning] PyPDFLoader failed for {path.name}: {e}")

        # Check if PyPDFLoader extracted any text
        has_text = any(doc.page_content.strip() for doc in docs)
        if not has_text:
            try:
                import pdfplumber

                with pdfplumber.open(str(path)) as pdf:
                    pdf_docs = []
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text() or ""
                        if text.strip():
                            pdf_docs.append(
                                Document(
                                    page_content=text,
                                    metadata={
                                        "source": str(path),
                                        "file_name": path.name,
                                        "file_type": "PDF",
                                        "page": i,
                                    },
                                )
                            )
                    if pdf_docs:
                        docs = pdf_docs
            except Exception as e:
                print(f"[Warning] pdfplumber fallback failed for {path.name}: {e}")

        for doc in docs:
            doc.metadata["file_name"] = path.name
            doc.metadata["file_type"] = "PDF"
        return docs

    def _load_docx(self, path: Path) -> List[Document]:
        """Load Microsoft Word document using python-docx."""
        try:
            import docx

            doc = docx.Document(str(path))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            text_content = "\n\n".join(full_text)
            metadata = {
                "source": str(path),
                "file_name": path.name,
                "file_type": "DOCX",
            }
            return [Document(page_content=text_content, metadata=metadata)]
        except Exception as e:
            # Fallback to Docx2txtLoader if available
            try:
                from langchain_community.document_loaders import Docx2txtLoader

                loader = Docx2txtLoader(str(path))
                docs = loader.load()
                for d in docs:
                    d.metadata["file_name"] = path.name
                    d.metadata["file_type"] = "DOCX"
                return docs
            except Exception:
                raise RuntimeError(f"Failed to load DOCX file {path.name}: {e}")

    def _load_txt(self, path: Path) -> List[Document]:
        """Load Plain Text or Markdown file."""
        loader = TextLoader(str(path), encoding="utf-8")
        docs = loader.load()
        for doc in docs:
            doc.metadata["file_name"] = path.name
            doc.metadata["file_type"] = "TXT"
        return docs

    def _load_excel_csv(self, path: Path) -> List[Document]:
        """Load Excel spreadsheet or CSV file, converting rows to structured documents."""
        docs = []
        ext = path.suffix.lower()

        if ext == ".csv":
            df_dict = {"Sheet1": pd.read_csv(str(path))}
        else:
            df_dict = pd.read_excel(str(path), sheet_name=None)

        for sheet_name, df in df_dict.items():
            if df.empty:
                continue
            # Format table content per sheet
            text_lines = [f"--- Sheet: {sheet_name} ---"]
            for idx, row in df.iterrows():
                row_str = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                text_lines.append(f"Row {idx + 1}: {row_str}")

            sheet_content = "\n".join(text_lines)
            metadata = {
                "source": str(path),
                "file_name": path.name,
                "file_type": "Excel/CSV",
                "sheet_name": sheet_name,
                "total_rows": len(df),
            }
            docs.append(Document(page_content=sheet_content, metadata=metadata))

        return docs

    def load_directory(self, directory_path: str) -> List[Document]:
        """Recursively scan a directory and load all supported documents."""
        dir_path = Path(directory_path)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Directory not found: {directory_path}")

        all_documents = []
        for path in dir_path.rglob("*"):
            if path.is_file() and path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    docs = self.load_file(str(path))
                    all_documents.extend(docs)
                except Exception as e:
                    print(f"[Warning] Failed to load {path.name}: {e}")

        return all_documents
