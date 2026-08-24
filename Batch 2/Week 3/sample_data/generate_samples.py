import os

def create_samples():
    sample_dir = os.path.dirname(os.path.abspath(__file__))

    # 1. TXT Sample
    txt_path = os.path.join(sample_dir, "sample_notes.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(
            "=== Alphatron Technologies RAG Guidelines ===\n"
            "Project Name: Advanced AI RAG Chatbot\n"
            "Version: 1.0.0\n"
            "Release Date: August 2026\n\n"
            "System Architecture Overview:\n"
            "The AI Chatbot uses Retrieval-Augmented Generation (RAG) to answer questions accurately based on custom document collections.\n"
            "It supports four core document types: PDF, DOCX, TXT, and Excel/CSV spreadsheets.\n\n"
            "Vector Storage:\n"
            "ChromaDB is used as the primary persistent vector database. Document chunks are embedded using sentence-transformers (all-MiniLM-L6-v2) or Google Gemini embeddings.\n\n"
            "Key Operational Metrics:\n"
            "- Target Retrieval Latency: < 500ms\n"
            "- Target Answer Accuracy: > 95%\n"
            "- Default Chunk Size: 1000 characters\n"
            "- Default Chunk Overlap: 200 characters\n"
        )

    # 2. CSV Sample
    csv_path = os.path.join(sample_dir, "sample_data.csv")
    with open(csv_path, "w", encoding="utf-8") as f:
        f.write(
            "EmployeeID,Name,Department,Role,Salary,Project\n"
            "E101,Alice Smith,Engineering,Lead RAG Architect,120000,Project RAG Chatbot\n"
            "E102,Bob Jones,Data Science,ML Engineer,95000,Project Vector\n"
            "E103,Carol White,Product,Product Manager,110000,Project RAG Chatbot\n"
            "E104,David Black,DevOps,Cloud Architect,105000,Project Infrastructure\n"
            "E105,Eva Green,Engineering,Backend Developer,88000,Project Vector\n"
        )

    # 3. DOCX Sample
    try:
        import docx
        doc = docx.Document()
        doc.add_heading("Alphatron Research Brief - DOCX Sample", 0)
        doc.add_paragraph(
            "This document presents research findings on vector indexing strategies using ChromaDB and LangChain."
        )
        doc.add_heading("Embedding Benchmark Results", level=1)
        doc.add_paragraph(
            "HuggingFace sentence-transformers achieved 94.2% accuracy on domain retrieval benchmarks, "
            "outperforming basic keyword TF-IDF searches by 38%."
        )
        doc.save(os.path.join(sample_dir, "sample_doc.docx"))
    except Exception as e:
        print(f"Skipping docx file creation: {e}")

    # 4. Excel Sample
    try:
        import pandas as pd
        df = pd.read_csv(csv_path)
        excel_path = os.path.join(sample_dir, "sample_data.xlsx")
        df.to_excel(excel_path, index=False)
    except Exception as e:
        print(f"Skipping excel file creation: {e}")

    # 5. PDF Sample
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        pdf_path = os.path.join(sample_dir, "sample_report.pdf")
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(100, 750, "=== Alphatron RAG Technical PDF Report ===")
        c.drawString(100, 730, "System Architecture: RAG AI Chatbot using ChromaDB Vector Store.")
        c.drawString(100, 710, "Supported Document Formats: PDF, DOCX, TXT, Excel/CSV.")
        c.drawString(100, 690, "Retrieval Engine: LangChain with sentence-transformers and Google Gemini.")
        c.save()
    except Exception as e:
        print(f"Skipping PDF file creation: {e}")

if __name__ == "__main__":
    create_samples()
    print("All sample datasets (PDF, DOCX, TXT, Excel/CSV) generated successfully!")
